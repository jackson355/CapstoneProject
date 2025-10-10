"""
Document Text Extraction Service
Extracts clean text content from DOCX files for AI analysis
"""

from typing import Dict, List, Any, Optional
import tempfile
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re


class DocumentTextExtractor:
    """
    Extracts structured text content from DOCX documents
    Focuses on text content only - no formatting, images, or binary data
    """

    def __init__(self):
        pass

    def extract_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text content from a DOCX file

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing structured text content
        """
        try:
            doc = Document(file_path)
            return self._extract_document_content(doc)
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")

    def extract_from_content(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text content from DOCX file bytes

        Args:
            file_content: Raw bytes of the DOCX file

        Returns:
            Dictionary containing structured text content
        """
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                doc = Document(temp_file_path)
                return self._extract_document_content(doc)
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)

        except Exception as e:
            raise Exception(f"Failed to extract text from document content: {str(e)}")

    def _extract_document_content(self, doc: Document) -> Dict[str, Any]:
        """
        Extract structured content from a Document object

        Args:
            doc: python-docx Document object

        Returns:
            Dictionary with extracted content
        """
        content = {
            "paragraphs": [],
            "headings": [],
            "tables": [],
            "metadata": {
                "total_paragraphs": 0,
                "total_tables": 0,
                "total_words": 0,
                "has_headers": False,
                "has_footers": False
            },
            "full_text": "",
            "structured_text": []
        }

        # Extract paragraphs and identify headings
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Determine if this is a heading based on style
            is_heading = self._is_heading_paragraph(paragraph)

            paragraph_data = {
                "text": text,
                "style": paragraph.style.name if paragraph.style else "Normal",
                "is_heading": is_heading,
                "alignment": self._get_alignment_text(paragraph.alignment),
                "runs_count": len(paragraph.runs)
            }

            content["paragraphs"].append(paragraph_data)
            content["structured_text"].append({
                "type": "heading" if is_heading else "paragraph",
                "content": text,
                "level": self._get_heading_level(paragraph.style.name) if is_heading else None
            })

            if is_heading:
                content["headings"].append({
                    "text": text,
                    "level": self._get_heading_level(paragraph.style.name),
                    "style": paragraph.style.name
                })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = self._extract_table_content(table, table_idx)
            content["tables"].append(table_data)

        # Extract headers and footers
        content["headers"] = self._extract_headers_footers(doc, "header")
        content["footers"] = self._extract_headers_footers(doc, "footer")

        # Build full text
        all_text_parts = []
        for item in content["structured_text"]:
            all_text_parts.append(item["content"])

        # Add table text
        for table in content["tables"]:
            for row in table["rows"]:
                for cell in row["cells"]:
                    if cell["text"].strip():
                        all_text_parts.append(cell["text"])

        content["full_text"] = "\n".join(all_text_parts)

        # Update metadata
        content["metadata"]["total_paragraphs"] = len(content["paragraphs"])
        content["metadata"]["total_tables"] = len(content["tables"])
        content["metadata"]["total_words"] = len(content["full_text"].split())
        content["metadata"]["has_headers"] = len(content["headers"]) > 0
        content["metadata"]["has_footers"] = len(content["footers"]) > 0

        return content

    def _is_heading_paragraph(self, paragraph) -> bool:
        """
        Determine if a paragraph is a heading based on style name
        """
        style_name = paragraph.style.name.lower()
        heading_indicators = [
            'heading', 'title', 'subtitle', 'header',
            'h1', 'h2', 'h3', 'h4', 'h5', 'h6'
        ]
        return any(indicator in style_name for indicator in heading_indicators)

    def _get_heading_level(self, style_name: str) -> Optional[int]:
        """
        Extract heading level from style name
        """
        if not style_name:
            return None

        style_lower = style_name.lower()

        # Look for numbered heading levels
        import re
        level_match = re.search(r'heading\s*(\d+)', style_lower)
        if level_match:
            return int(level_match.group(1))

        # Look for h1, h2, etc.
        h_match = re.search(r'h(\d+)', style_lower)
        if h_match:
            return int(h_match.group(1))

        # Default heading levels
        if 'title' in style_lower:
            return 1
        elif 'subtitle' in style_lower:
            return 2

        return 1

    def _get_alignment_text(self, alignment) -> str:
        """
        Convert alignment enum to readable text
        """
        if alignment is None:
            return "left"

        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify"
        }

        return alignment_map.get(alignment, "left")

    def _extract_table_content(self, table, table_idx: int) -> Dict[str, Any]:
        """
        Extract content from a table
        """
        table_data = {
            "index": table_idx,
            "rows_count": len(table.rows),
            "cols_count": len(table.columns) if table.rows else 0,
            "rows": [],
            "text_content": []
        }

        for row_idx, row in enumerate(table.rows):
            row_data = {
                "index": row_idx,
                "cells": []
            }

            row_texts = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_data = {
                    "index": cell_idx,
                    "text": cell_text
                }
                row_data["cells"].append(cell_data)
                if cell_text:
                    row_texts.append(cell_text)

            table_data["rows"].append(row_data)
            if row_texts:
                table_data["text_content"].append(" | ".join(row_texts))

        return table_data

    def _extract_headers_footers(self, doc: Document, section_type: str) -> List[str]:
        """
        Extract text from headers or footers
        """
        content = []

        for section in doc.sections:
            if section_type == "header":
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content.append(text)
            elif section_type == "footer":
                footer = section.footer
                if footer:
                    for paragraph in footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content.append(text)

        return content

    def get_clean_text_for_ai(self, extracted_content: Dict[str, Any]) -> str:
        """
        Get clean, AI-friendly text from extracted content

        Args:
            extracted_content: Output from extract_from_file or extract_from_content

        Returns:
            Clean text suitable for AI analysis
        """
        text_parts = []

        # Add document structure
        text_parts.append("=== DOCUMENT CONTENT ===\n")

        # Add structured content in order
        for item in extracted_content["structured_text"]:
            if item["type"] == "heading":
                level_prefix = "#" * (item["level"] or 1)
                text_parts.append(f"{level_prefix} {item['content']}")
            else:
                text_parts.append(item["content"])
            text_parts.append("")  # Add blank line

        # Add tables if present
        if extracted_content["tables"]:
            text_parts.append("\n=== TABLES ===\n")
            for table_idx, table in enumerate(extracted_content["tables"]):
                text_parts.append(f"Table {table_idx + 1}:")
                for row_text in table["text_content"]:
                    text_parts.append(row_text)
                text_parts.append("")

        # Add headers/footers if present
        if extracted_content["headers"]:
            text_parts.append("\n=== HEADERS ===\n")
            text_parts.extend(extracted_content["headers"])

        if extracted_content["footers"]:
            text_parts.append("\n=== FOOTERS ===\n")
            text_parts.extend(extracted_content["footers"])

        return "\n".join(text_parts)

    def extract_potential_variables(self, text: str) -> List[Dict[str, Any]]:
        """
        Pre-analyze text to identify potential variables before sending to AI

        Args:
            text: Clean text content

        Returns:
            List of potential variable locations and patterns
        """
        variables = []

        # Common patterns that suggest variables
        patterns = [
            # Bracketed placeholders
            (r'\[([^\]]+)\]', 'bracketed'),
            (r'\{\{([^}]+)\}\}', 'mustache'),
            (r'\{([^}]+)\}', 'curly'),

            # XXX patterns
            (r'\bXXX+\b', 'xxx_pattern'),
            (r'_{3,}', 'underlines'),

            # Common business data patterns
            (r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', 'date'),
            (r'\$[\d,]+\.?\d*', 'currency'),
            (r'\b\d+\.\d{2}\b', 'decimal'),
            (r'\b[A-Z]{2,}\s+[A-Z]{2,}', 'company_caps'),

            # Email and contact patterns
            (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'email'),
            (r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', 'phone'),

            # Address patterns
            (r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Lane|Ln|Drive|Dr|Boulevard|Blvd)\b', 'address'),
        ]

        for pattern, var_type in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                variables.append({
                    'text': match.group(0),
                    'type': var_type,
                    'start': match.start(),
                    'end': match.end(),
                    'context': text[max(0, match.start()-50):match.end()+50].strip()
                })

        return variables


# Create a global instance
document_extractor = DocumentTextExtractor()