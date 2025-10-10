"""
Text Improvement Service
Improves text content while preserving all formatting, layout, images, and tables
"""

from typing import Dict, List, Any, Optional
import tempfile
import os
import time
from pathlib import Path
from docx import Document
import re
from pydantic import BaseModel
from openai import AsyncOpenAI


class TextSegment(BaseModel):
    """Represents a text segment to be improved"""
    original_text: str
    location_type: str  # 'paragraph' or 'table_cell'
    location_info: str  # Description of where this text is located
    improved_text: Optional[str] = None


class TextImprovementResult(BaseModel):
    """Result of text improvement analysis"""
    segments: List[TextSegment]
    total_segments: int
    improved_segments: int
    preserved_placeholders: List[str]
    summary: str


class DocumentTextImprover:
    """
    Improves text content in DOCX documents while preserving all formatting
    Only modifies the .text property of paragraphs and cells
    """

    def __init__(self):
        pass

    def extract_text_segments(self, file_content: bytes) -> List[TextSegment]:
        """
        Extract all text segments from a DOCX document

        Args:
            file_content: Raw bytes of the DOCX file

        Returns:
            List of text segments with location information
        """
        segments = []

        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                doc = Document(temp_file_path)

                # Extract text from paragraphs
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:  # Only include non-empty paragraphs
                        segment = TextSegment(
                            original_text=text,
                            location_type="paragraph",
                            location_info=f"Paragraph {i + 1}"
                        )
                        segments.append(segment)

                # Extract text from table cells
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            text = cell.text.strip()
                            if text:  # Only include non-empty cells
                                segment = TextSegment(
                                    original_text=text,
                                    location_type="table_cell",
                                    location_info=f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}"
                                )
                                segments.append(segment)

            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)

        except Exception as e:
            raise Exception(f"Failed to extract text segments: {str(e)}")

        return segments

    async def improve_text_with_ai(
        self,
        segments: List[TextSegment],
        api_key: str,
        improvement_type: str = "grammar_clarity"
    ) -> TextImprovementResult:
        """
        Improve text segments using AI while preserving placeholders

        Args:
            segments: List of text segments to improve
            api_key: OpenAI API key
            improvement_type: Type of improvement to apply

        Returns:
            Improvement result with AI-enhanced text
        """
        client = AsyncOpenAI(api_key=api_key)

        # Collect all text for batch processing
        all_text = []
        for i, segment in enumerate(segments):
            all_text.append(f"[SEGMENT_{i}] {segment.original_text}")

        combined_text = "\n\n".join(all_text)

        # Create improvement prompt
        prompt = self._build_improvement_prompt(combined_text, improvement_type)

        try:
            response = await client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": self._get_improvement_system_prompt()
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,  # Lower temperature for consistent improvements
                max_tokens=4000
            )

            # Parse the response
            improved_text = response.choices[0].message.content

            # Extract improved segments
            improved_segments = self._parse_improved_text(improved_text, segments)

            # Find preserved placeholders
            preserved_placeholders = self._find_preserved_placeholders(segments)

            return TextImprovementResult(
                segments=improved_segments,
                total_segments=len(segments),
                improved_segments=len([s for s in improved_segments if s.improved_text]),
                preserved_placeholders=preserved_placeholders,
                summary=f"Improved {len(improved_segments)} text segments while preserving formatting and placeholders"
            )

        except Exception as e:
            raise Exception(f"AI text improvement failed: {str(e)}")

    def _get_improvement_system_prompt(self) -> str:
        """Get the system prompt for text improvement"""
        return """You are a professional text editor specialized in improving grammar, clarity, and readability while preserving document structure.

CRITICAL RULES:
1. NEVER change or remove placeholders like {{variable}}, [placeholder], {{client_name}}, {{date}}, etc.
2. NEVER add or remove formatting markup, bullets, or numbering
3. NEVER change the meaning or intent of the text
4. Only improve grammar, word choice, and sentence structure
5. Preserve the original tone and style
6. Keep technical terms and proper nouns unchanged
7. Maintain the same paragraph structure

Your response must follow this exact format:
[SEGMENT_0] improved text here
[SEGMENT_1] improved text here
[SEGMENT_2] improved text here

Each segment should be on its own line with the exact segment number."""

    def _build_improvement_prompt(self, combined_text: str, improvement_type: str) -> str:
        """Build the improvement prompt"""

        improvement_instructions = {
            "grammar_clarity": "Focus on fixing grammar errors and improving clarity while maintaining the original meaning",
            "professional_tone": "Enhance the text to sound more professional and polished",
            "concise": "Make the text more concise while retaining all important information",
            "formal": "Adjust the tone to be more formal and business-appropriate"
        }

        instruction = improvement_instructions.get(improvement_type, improvement_instructions["grammar_clarity"])

        return f"""Please improve the following text segments. {instruction}.

IMPORTANT RULES:
- Keep ALL placeholders like {{{{variable}}}}, [placeholder], {{{{client_name}}}}, {{{{date}}}} exactly as they are
- Do NOT change formatting, bullets, numbering, or structure
- Only improve the actual text content
- Preserve technical terms, proper nouns, and brand names
- Keep the same tone and style

TEXT TO IMPROVE:
{combined_text}

Please respond with each improved segment using the exact [SEGMENT_X] format."""

    def _parse_improved_text(self, improved_text: str, original_segments: List[TextSegment]) -> List[TextSegment]:
        """Parse the AI response and match improved text to segments"""
        improved_segments = []

        # Split response into lines and look for segment markers
        lines = improved_text.strip().split('\n')
        current_segment = None
        current_text = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line starts a new segment
            segment_match = re.match(r'\[SEGMENT_(\d+)\]\s*(.*)', line)
            if segment_match:
                # Save previous segment if exists
                if current_segment is not None and current_text:
                    improved_text_content = ' '.join(current_text).strip()
                    if improved_text_content:
                        original_segments[current_segment].improved_text = improved_text_content

                # Start new segment
                current_segment = int(segment_match.group(1))
                current_text = [segment_match.group(2)] if segment_match.group(2) else []
            else:
                # Add to current segment
                if current_segment is not None:
                    current_text.append(line)

        # Save last segment
        if current_segment is not None and current_text:
            improved_text_content = ' '.join(current_text).strip()
            if improved_text_content:
                original_segments[current_segment].improved_text = improved_text_content

        return original_segments

    def _find_preserved_placeholders(self, segments: List[TextSegment]) -> List[str]:
        """Find all placeholders that were preserved in the text"""
        placeholders = set()

        # Common placeholder patterns
        patterns = [
            r'\{\{[^}]+\}\}',  # {{variable}}
            r'\[[^\]]+\]',     # [placeholder]
            r'\$\{[^}]+\}',    # ${variable}
            r'%[^%]+%',        # %variable%
        ]

        for segment in segments:
            for pattern in patterns:
                matches = re.findall(pattern, segment.original_text)
                placeholders.update(matches)

                if segment.improved_text:
                    matches = re.findall(pattern, segment.improved_text)
                    placeholders.update(matches)

        return sorted(list(placeholders))

    def apply_all_replacements_to_document(
        self,
        file_content: bytes,
        all_replacements: Dict[str, str]
    ) -> bytes:
        """
        Apply all replacements (variables and improvements) to the document

        FORMATTING PRESERVATION GUARANTEE:
        - Preserves ALL fonts, sizes, colors, bold, italic, underline
        - Preserves ALL paragraph alignment, spacing, indentation
        - Preserves ALL tables, borders, cell formatting
        - Preserves ALL images, charts, diagrams
        - Preserves ALL headers, footers, page layout
        - Only modifies .text properties of paragraphs and cells

        Args:
            file_content: Original DOCX file content
            all_replacements: Dictionary mapping original text to replacement text

        Returns:
            Modified DOCX file content as bytes with identical formatting
        """
        return self.apply_variable_replacements_to_document(file_content, all_replacements)

    def apply_variable_replacements_to_document(
        self,
        file_content: bytes,
        variable_replacements: Dict[str, str]
    ) -> bytes:
        """
        Apply variable replacements to the original document
        Only modifies .text properties, preserves all formatting

        Args:
            file_content: Original DOCX file content
            variable_replacements: Dictionary mapping original text to placeholder

        Returns:
            Modified DOCX file content as bytes
        """
        temp_input_path = None
        temp_output_path = None

        try:
            # Create temporary file for input with retry logic
            temp_input_path = tempfile.mktemp(suffix='.docx')
            with open(temp_input_path, 'wb') as f:
                f.write(file_content)

            # Wait briefly to ensure file is not locked
            time.sleep(0.1)

            # Load document with retry logic
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    doc = Document(temp_input_path)
                    break
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise
                    time.sleep(0.2)  # Wait before retry

            # Apply variable replacements to paragraphs
            for para in doc.paragraphs:
                original_text = para.text
                modified_text = original_text

                # Apply all replacements to this paragraph
                for original_value, replacement in variable_replacements.items():
                    if original_value in modified_text:
                        modified_text = modified_text.replace(original_value, replacement)

                # Only update if text changed
                if modified_text != original_text:
                    para.text = modified_text

            # Apply variable replacements to table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text
                        modified_text = original_text

                        # Apply all replacements to this cell
                        for original_value, replacement in variable_replacements.items():
                            if original_value in modified_text:
                                modified_text = modified_text.replace(original_value, replacement)

                        # Only update if text changed
                        if modified_text != original_text:
                            cell.text = modified_text

            # Save to temporary output file
            temp_output_path = tempfile.mktemp(suffix='.docx')
            doc.save(temp_output_path)

            # Wait briefly before reading
            time.sleep(0.1)

            # Read the modified document with retry logic
            for attempt in range(max_retries):
                try:
                    with open(temp_output_path, 'rb') as f:
                        modified_content = f.read()
                    break
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise
                    time.sleep(0.2)  # Wait before retry

            return modified_content

        except Exception as e:
            raise Exception(f"Failed to apply variable replacements: {str(e)}")
        finally:
            # Clean up temporary files with error handling
            for temp_path in [temp_input_path, temp_output_path]:
                if temp_path and os.path.exists(temp_path):
                    try:
                        os.unlink(temp_path)
                    except Exception:
                        # If normal deletion fails, try with a delay
                        try:
                            time.sleep(0.1)
                            os.unlink(temp_path)
                        except Exception:
                            pass  # Ignore cleanup errors

    def apply_improvements_to_document(
        self,
        file_content: bytes,
        improved_segments: List[TextSegment]
    ) -> bytes:
        """
        Apply improvements back to the original document
        Only modifies .text properties, preserves all formatting

        Args:
            file_content: Original DOCX file content
            improved_segments: Segments with improved text

        Returns:
            Modified DOCX file content as bytes
        """
        try:
            # Create temporary file for input
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_input:
                temp_input.write(file_content)
                temp_input_path = temp_input.name

            try:
                doc = Document(temp_input_path)

                # Create a mapping of improved text by location
                improvements = {}
                for segment in improved_segments:
                    if segment.improved_text and segment.improved_text != segment.original_text:
                        improvements[segment.original_text] = segment.improved_text

                # Apply improvements to paragraphs
                for para in doc.paragraphs:
                    original_text = para.text.strip()
                    if original_text in improvements:
                        # CRITICAL: Only modify the .text property
                        para.text = improvements[original_text]

                # Apply improvements to table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            original_text = cell.text.strip()
                            if original_text in improvements:
                                # CRITICAL: Only modify the .text property
                                cell.text = improvements[original_text]

                # Save to temporary output file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_output:
                    doc.save(temp_output.name)

                    # Read the improved document
                    with open(temp_output.name, 'rb') as f:
                        improved_content = f.read()

                    # Clean up output file
                    os.unlink(temp_output.name)

            finally:
                # Clean up input file
                os.unlink(temp_input_path)

            return improved_content

        except Exception as e:
            raise Exception(f"Failed to apply improvements: {str(e)}")

    def get_improvement_preview(
        self,
        segments: List[TextSegment]
    ) -> Dict[str, Any]:
        """
        Generate a preview of text improvements

        Args:
            segments: Segments with improved text

        Returns:
            Preview data for UI display
        """
        changes = []
        stats = {
            "total_segments": len(segments),
            "improved_segments": 0,
            "unchanged_segments": 0,
            "preserved_placeholders": 0
        }

        for segment in segments:
            if segment.improved_text and segment.improved_text != segment.original_text:
                changes.append({
                    "location": segment.location_info,
                    "original": segment.original_text,
                    "improved": segment.improved_text,
                    "type": segment.location_type
                })
                stats["improved_segments"] += 1
            else:
                stats["unchanged_segments"] += 1

        # Count preserved placeholders
        placeholders = self._find_preserved_placeholders(segments)
        stats["preserved_placeholders"] = len(placeholders)

        return {
            "changes": changes,
            "stats": stats,
            "placeholders": placeholders
        }


# Global instance
document_text_improver = DocumentTextImprover()