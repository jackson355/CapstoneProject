"""
Quotation Placeholder Filler Service

This service fills template placeholders with client data for quotations.
"""

from docx import Document
from io import BytesIO
import re
from typing import Dict, List, Tuple
from datetime import datetime


class QuotationFillerService:
    """Service to fill quotation templates with client data"""

    def __init__(self):
        self.placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')

    def fill_template_with_client_data(
        self,
        template_content: bytes,
        client_data: Dict,
        contact_data: Dict,
        company_data: Dict = None
    ) -> Tuple[bytes, List[str]]:
        """
        Fill a DOCX template with client and contact information.

        Args:
            template_content: Bytes of the DOCX template
            client_data: Dictionary with client information
            contact_data: Dictionary with contact person information
            company_data: Dictionary with user's company information (optional)

        Returns:
            Tuple of (filled_content_bytes, unfilled_placeholders)
        """

        # Load the document
        doc = Document(BytesIO(template_content))

        # Create mapping of placeholders to values
        placeholder_map = self._create_placeholder_map(client_data, contact_data, company_data or {})

        # Track unfilled placeholders
        unfilled = set()

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            unfilled.update(self._fill_paragraph(paragraph, placeholder_map))

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        unfilled.update(self._fill_paragraph(paragraph, placeholder_map))

        # Replace placeholders in headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                unfilled.update(self._fill_paragraph(paragraph, placeholder_map))

        # Replace placeholders in footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                unfilled.update(self._fill_paragraph(paragraph, placeholder_map))

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output.read(), list(unfilled)

    def _create_placeholder_map(self, client_data: Dict, contact_data: Dict, company_data: Dict) -> Dict[str, str]:
        """
        Create a mapping of placeholder names to their values.

        Standard placeholders:
        - {{client_company_name}} - Company name
        - {{client_uen}} - UEN number
        - {{client_industry}} - Industry
        - {{client_address}} - Address
        - {{client_postal_code}} - Postal code
        - {{contact_name}} - Contact person name
        - {{contact_phone}} - Contact person phone
        - {{contact_email}} - Contact person email
        - {{my_company_name}} - User's company name
        - {{my_company_email}} - User's company email
        - {{my_company_phone}} - User's company phone
        - {{my_company_address}} - User's company address
        - {{my_company_website}} - User's company website
        - {{current_date}} - Current date
        - {{quotation_date}} - Same as current date
        """

        # Create mapping
        mapping = {}

        # Client information
        mapping['client_company_name'] = str(client_data.get('company_name') or '')
        mapping['client_uen'] = str(client_data.get('uen') or '')
        mapping['client_industry'] = str(client_data.get('industry') or '')
        mapping['client_address'] = str(client_data.get('address') or '')
        mapping['client_postal_code'] = str(client_data.get('postal_code') or '')

        # Contact information
        mapping['contact_name'] = str(contact_data.get('name') or '')
        mapping['contact_phone'] = str(contact_data.get('phone') or '')
        mapping['contact_email'] = str(contact_data.get('email') or '')

        # User's company information
        mapping['my_company_name'] = str(company_data.get('name') or '')
        mapping['my_company_email'] = str(company_data.get('email') or '')
        mapping['my_company_phone'] = str(company_data.get('phone') or '')
        mapping['my_company_address'] = str(company_data.get('address') or '')
        mapping['my_company_website'] = str(company_data.get('website') or '')

        # Date placeholders
        current_date = datetime.now().strftime('%d/%m/%Y')
        mapping['current_date'] = current_date
        mapping['quotation_date'] = current_date
        mapping['date'] = current_date

        return mapping

    def _fill_paragraph(self, paragraph, placeholder_map: Dict[str, str]) -> set:
        """
        Fill placeholders in a paragraph while preserving formatting.
        Returns set of unfilled placeholders.
        """
        unfilled = set()

        # Get the full text to find placeholders
        full_text = paragraph.text

        # Find all placeholders
        placeholders = self.placeholder_pattern.findall(full_text)

        if not placeholders:
            return unfilled

        # For each run in the paragraph, replace placeholders
        for run in paragraph.runs:
            text = run.text

            for match in self.placeholder_pattern.finditer(text):
                placeholder_name = match.group(1).strip()
                placeholder_full = match.group(0)  # {{placeholder}}

                # Check if we have a value for this placeholder
                if placeholder_name in placeholder_map:
                    value = placeholder_map[placeholder_name]
                    if value:  # Only replace if we have a non-empty value
                        text = text.replace(placeholder_full, value)
                    else:
                        unfilled.add(placeholder_name)
                else:
                    unfilled.add(placeholder_name)

            run.text = text

        return unfilled

    def extract_all_placeholders(self, template_content: bytes) -> List[str]:
        """
        Extract all unique placeholders from a template.

        Args:
            template_content: Bytes of the DOCX template

        Returns:
            List of unique placeholder names (without {{ }})
        """
        doc = Document(BytesIO(template_content))
        placeholders = set()

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            placeholders.update(self.placeholder_pattern.findall(paragraph.text))

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(self.placeholder_pattern.findall(paragraph.text))

        # Extract from headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                placeholders.update(self.placeholder_pattern.findall(paragraph.text))

        # Extract from footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                placeholders.update(self.placeholder_pattern.findall(paragraph.text))

        return sorted([p.strip() for p in placeholders])


# Global instance
quotation_filler = QuotationFillerService()
