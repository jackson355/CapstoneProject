from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form, Request
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
import json
import uuid
import os
import tempfile
from pathlib import Path

from app.db.session import get_db
from app.models import Template, User, Role, ActivityLog, Quotation, Invoice
from app.schemas import TemplateCreate, TemplateUpdate, TemplateOut, TemplateListItem, PaginatedTemplatesResponse, UserOut
from app.api.auth import get_current_user
from app.services.file_storage import file_storage
from app.services.document_extractor import document_extractor
from app.services.ai_template_analyzer import ai_template_analyzer
from app.services.text_improver import document_text_improver
from app.services.quotation_filler import quotation_filler

router = APIRouter(prefix="/templates", tags=["templates"])

# OnlyOffice configuration from environment variables
BACKEND_CALLBACK_URL = os.getenv("BACKEND_CALLBACK_URL", "http://host.docker.internal:8000")

# RBAC dependency for admin/superadmin access
def require_admin_or_superadmin(current_user: UserOut = Depends(get_current_user), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == current_user.id).first()
    if not user:
        raise HTTPException(status_code=403, detail="User not found")
    role = db.query(Role).filter(Role.id == user.role_id).first()
    if not role or role.name not in ["admin", "superadmin"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    return current_user

@router.get("/", response_model=PaginatedTemplatesResponse)
def list_templates(
    page: int = Query(default=0, ge=0),
    per_page: int = Query(default=10, ge=1, le=100),
    search: Optional[str] = Query(default=None, description="Search by name or description"),
    template_type: Optional[str] = Query(default=None),
    status: Optional[str] = Query(default=None),
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    query = db.query(Template)

    if search:
        term = f"%{search}%"
        query = query.filter(
            (Template.name.ilike(term)) |
            (Template.description.ilike(term))
        )

    if template_type:
        query = query.filter(Template.template_type == template_type)

    if status:
        query = query.filter(Template.status == status)

    total = query.count()
    templates = query.order_by(Template.updated_at.desc()).offset(page * per_page).limit(per_page).all()

    print(f"Fetched {len(templates)} templates from database")  # Debug log
    for template in templates:
        print(f"Template {template.id}: content = {template.content}")  # Debug log

    return PaginatedTemplatesResponse(
        templates=templates,
        total=total,
        page=page,
        per_page=per_page
    )

@router.post("/", response_model=TemplateOut)
def create_template(
    template_in: TemplateCreate,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    template_data = template_in.dict()
    template_data['created_by'] = current_user.id

    print(f"Creating template with data: {template_data}")  # Debug log

    template = Template(**template_data)
    db.add(template)
    db.commit()
    db.refresh(template)

    # Create a blank DOCX file for OnlyOffice editing
    try:
        from docx import Document
        import tempfile

        # Create a blank document
        doc = Document()

        # Add some initial content
        if template.content and isinstance(template.content, dict) and template.content.get('html'):
            # Simple HTML to text conversion for initial content
            import re
            html_content = template.content.get('html', '')
            text_content = re.sub('<[^<]+?>', '', html_content)
            if text_content.strip():
                doc.add_paragraph(text_content)
            else:
                doc.add_paragraph(f"Template: {template.name}")
                doc.add_paragraph("Start editing your document here...")
        else:
            doc.add_paragraph(f"Template: {template.name}")
            doc.add_paragraph("Start editing your document here...")

        # Save the DOCX file using file_storage
        temp_dir = Path(tempfile.gettempdir()) / "template_creation"
        temp_dir.mkdir(exist_ok=True)
        temp_file = temp_dir / f"temp_{template.id}.docx"

        doc.save(str(temp_file))

        # Read the file and save it using file_storage
        with open(temp_file, 'rb') as f:
            file_content = f.read()

        # Save using the file storage system
        file_metadata = file_storage.save_docx_content(file_content, template.id)

        # Update template with file information
        template.file_path = f"uploads/templates/template_{template.id}.docx"
        template.file_name = f"{template.name}.docx"
        template.file_size = len(file_content)
        template.mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Update content to include file metadata
        template.content = {
            "type": "docx",
            "file_url": f"/templates/document/{template.id}",
            "file_name": template.file_name,
            "file_size": template.file_size,
            "mime_type": template.mime_type
        }

        # Clean up temp file
        temp_file.unlink()

        db.commit()
        db.refresh(template)

        print(f"Created DOCX file for template: {template.id}")

    except Exception as e:
        print(f"Error creating DOCX file for template {template.id}: {e}")
        # Continue without DOCX file - template will still work with HTML content

    print(f"Created template: {template.id}, content: {template.content}")  # Debug log

    # Log activity
    db.add(ActivityLog(
        action="template.create",
        actor_user_id=current_user.id,
        target_type="template",
        target_id=template.id,
        message=f"Template '{template.name}' created",
        log_metadata={"template_type": template.template_type}
    ))
    db.commit()

    return template

@router.post("/upload-docx", response_model=TemplateOut)
async def upload_docx_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    template_type: str = Form("document"),
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """Upload a DOCX file and create a new template"""

    # Create template record first
    template_data = {
        "name": name,
        "description": description,
        "template_type": template_type,
        "content": {"type": "docx"},  # Placeholder, will be updated
        "created_by": current_user.id,
        "status": "saved"
    }

    template = Template(**template_data)
    db.add(template)
    db.commit()
    db.refresh(template)

    try:
        # Save the DOCX file
        file_metadata = await file_storage.save_docx_file(file, template.id)

        # Update template with file information
        template.file_path = file_metadata["file_path"]
        template.file_name = file_metadata["file_name"]
        template.file_size = file_metadata["file_size"]
        template.mime_type = file_metadata["mime_type"]

        # Update content to include file metadata
        template.content = {
            "type": "docx",
            "file_url": f"/templates/document/{template.id}",
            "file_name": file_metadata["file_name"],
            "file_size": file_metadata["file_size"],
            "mime_type": file_metadata["mime_type"]
        }

        db.commit()
        db.refresh(template)

        # Log activity
        db.add(ActivityLog(
            action="template.upload_docx",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template.id,
            message=f"DOCX template '{template.name}' uploaded",
            log_metadata={"file_name": file_metadata["file_name"], "file_size": file_metadata["file_size"]}
        ))
        db.commit()

        return template

    except Exception as e:
        # Clean up template record if file save failed
        db.delete(template)
        db.commit()
        raise HTTPException(status_code=500, detail=f"Failed to upload template: {str(e)}")

@router.get("/{template_id}", response_model=TemplateOut)
def get_template(
    template_id: int,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return template

@router.put("/{template_id}", response_model=TemplateOut)
def update_template(
    template_id: int,
    template_in: TemplateUpdate,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    update_data = template_in.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(template, key, value)

    db.commit()
    db.refresh(template)

    # Log activity
    db.add(ActivityLog(
        action="template.update",
        actor_user_id=current_user.id,
        target_type="template",
        target_id=template.id,
        message=f"Template '{template.name}' updated",
        log_metadata={"template_type": template.template_type}
    ))
    db.commit()

    return template

@router.delete("/{template_id}")
def delete_template(
    template_id: int,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    # Check if template is being used by any quotations
    quotations_using_template = db.query(Quotation).filter(Quotation.template_id == template_id).all()

    # Check if template is being used by any invoices
    invoices_using_template = db.query(Invoice).filter(Invoice.template_id == template_id).all()

    # If template is in use, prevent deletion
    if quotations_using_template or invoices_using_template:
        error_details = {
            "message": "Cannot delete template because it is currently being used",
            "template_name": template.name,
            "usage": {
                "quotations": [
                    {
                        "id": q.id,
                        "quotation_number": q.quotation_number,
                        "created_at": q.created_at.isoformat()
                    }
                    for q in quotations_using_template
                ],
                "invoices": [
                    {
                        "id": inv.id,
                        "invoice_number": inv.invoice_number,
                        "created_at": inv.created_at.isoformat()
                    }
                    for inv in invoices_using_template
                ]
            },
            "total_quotations": len(quotations_using_template),
            "total_invoices": len(invoices_using_template)
        }

        raise HTTPException(
            status_code=400,
            detail=error_details
        )

    template_name = template.name
    template_type = template.template_type

    # Delete the associated DOCX file
    try:
        file_storage.delete_docx_file(template_id)
    except Exception as e:
        print(f"Warning: Could not delete template file {template_id}: {str(e)}")

    db.delete(template)
    db.commit()

    # Log activity
    db.add(ActivityLog(
        action="template.delete",
        actor_user_id=current_user.id,
        target_type="template",
        target_id=template_id,
        message=f"Template '{template_name}' deleted",
        log_metadata={"template_type": template_type}
    ))
    db.commit()

    return {"detail": "Template deleted"}





# OnlyOffice Integration Endpoints
@router.api_route("/document/{template_id}", methods=["GET", "HEAD"])
def get_document_for_editing(
    template_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    Serve DOCX file for OnlyOffice editor - supports both GET and HEAD
    """
    print(f"[DEBUG] OnlyOffice Document Request - ID: {template_id}, Method: {request.method}")

    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        print(f"[ERROR] Template {template_id} not found")
        raise HTTPException(status_code=404, detail="Template not found")

    print(f"[SUCCESS] Template found: {template.name}")
    print(f"[INFO] Template file_path: {template.file_path}")
    print(f"[INFO] Template content type: {type(template.content)}")
    if template.content:
        print(f"[DEBUG] Content keys: {list(template.content.keys()) if isinstance(template.content, dict) else 'Not a dict'}")

    # Check if template has DOCX file
    if template.file_path and file_storage.file_exists(template_id):
        return file_storage.get_docx_file(template_id)

    # For legacy templates with only HTML content, create a DOCX file
    if template.content and isinstance(template.content, dict) and template.content.get('html'):
        from docx import Document
        import re

        doc = Document()

        # Simple HTML to text conversion
        html_content = template.content.get('html', '')
        text_content = re.sub('<[^<]+?>', '', html_content)
        doc.add_paragraph(text_content)

        # Save to temporary file and return
        temp_dir = Path(tempfile.gettempdir()) / "onlyoffice_docs"
        temp_dir.mkdir(exist_ok=True)

        file_path = temp_dir / f"template_{template_id}.docx"
        doc.save(str(file_path))

        return FileResponse(
            path=str(file_path),
            filename=f"{template.name}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    raise HTTPException(status_code=404, detail="No document file available for this template")

@router.post("/save/{template_id}")
async def save_document_callback(
    template_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    OnlyOffice callback endpoint for saving edited documents
    """
    print(f"[INFO] OnlyOffice save callback for template {template_id}")

    try:
        # Parse the JSON body from OnlyOffice
        body = await request.json()
        print(f"[DEBUG] OnlyOffice callback data: {body}")

        status = body.get("status")

        # Status codes: 1=editing, 2=ready for saving, 3=saving error, 4=closed with no changes
        if status == 2:  # Ready for saving
            download_url = body.get("url")
            if not download_url:
                print("[ERROR] No download URL provided in callback")
                return {"error": 1, "message": "No download URL provided"}

            template = db.query(Template).filter(Template.id == template_id).first()
            if not template:
                print(f"[ERROR] Template {template_id} not found")
                return {"error": 1, "message": "Template not found"}

            # Download the saved document from OnlyOffice
            print(f"[INFO] Downloading saved document from: {download_url}")
            import aiohttp
            async with aiohttp.ClientSession() as session:
                async with session.get(download_url) as response:
                    if response.status == 200:
                        file_content = await response.read()

                        # Save the updated document
                        file_storage.save_docx_content(file_content, template_id)

                        # Update template metadata
                        template.file_size = len(file_content)
                        template.updated_at = datetime.utcnow()

                        # Clear any cached HTML preview
                        if isinstance(template.content, dict) and "html_preview" in template.content:
                            del template.content["html_preview"]

                        db.commit()

                        print(f"[SUCCESS] Successfully saved template {template_id}")
                        return {"error": 0}
                    else:
                        print(f"[ERROR] Failed to download document: {response.status}")
                        return {"error": 1, "message": f"Failed to download document: {response.status}"}

        elif status == 1:  # Still editing
            print("[DEBUG] Document still being edited")
            return {"error": 0}
        elif status == 4:  # Closed without changes
            print("[INFO] Document closed without changes")
            return {"error": 0}
        else:
            print(f"[DEBUG] Unknown status: {status}")
            return {"error": 0}

    except Exception as e:
        print(f"[ERROR] Error in save callback: {e}")
        import traceback
        traceback.print_exc()
        return {"error": 1, "message": str(e)}


@router.get("/onlyoffice-config/{template_id}")
def get_onlyoffice_config(
    template_id: int,
    current_user: UserOut = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get OnlyOffice configuration for a specific template
    """
    print(f"[CONFIG] OnlyOffice Config Request - ID: {template_id}, User: {current_user.name}")

    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        print(f"[ERROR] Template {template_id} not found for config")
        raise HTTPException(status_code=404, detail="Template not found")

    config = {
        "document": {
            "fileType": "docx",
            "key": f"{template_id}_{int(datetime.utcnow().timestamp())}",
            "title": template.name,
            "url": f"{BACKEND_CALLBACK_URL}/templates/document/{template_id}",
            "permissions": {
                "edit": True,
                "download": True,
                "print": True,
                "review": True,
                "comment": True
            }
        },
        "documentType": "word",
        "editorConfig": {
            "callbackUrl": f"{BACKEND_CALLBACK_URL}/templates/save/{template_id}",
            "mode": "edit",
            "lang": "en",
            "user": {
                "id": str(current_user.id),
                "name": current_user.name,
                "group": "editors"
            },
            "customization": {
                "autosave": True,
                "forcesave": True
            }
        },
        "height": "100%",
        "width": "100%"
    }

    print(f"[INFO] Returning OnlyOffice config for template {template_id}:")
    print(f"   Document URL: {config['document']['url']}")
    print(f"   Document Key: {config['document']['key']}")
    print(f"   Document Type: {config['documentType']}")
    print(f"   Callback URL: {config['editorConfig']['callbackUrl']}")

    return config



# AI Template Conversion Endpoints

@router.post("/ai-analyze/{template_id}")
async def analyze_template_with_ai(
    template_id: int,
    openai_api_key: str = Form(...),
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """
    Analyze a template with AI to suggest variable improvements
    """
    print(f"[AI ANALYZE] Starting AI analysis for template {template_id}")

    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    try:
        # Extract text from the template's DOCX file
        if not template.file_path or not file_storage.file_exists(template_id):
            raise HTTPException(status_code=400, detail="No DOCX file available for analysis")

        # Get the DOCX file content
        file_content = file_storage.get_docx_content(template_id)

        # Extract structured text
        extracted_content = document_extractor.extract_from_content(file_content)
        clean_text = document_extractor.get_clean_text_for_ai(extracted_content)

        print(f"[AI ANALYZE] Extracted {len(clean_text)} characters of text")

        # Prepare metadata
        metadata = {
            "template_name": template.name,
            "template_type": template.template_type,
            "file_name": template.file_name,
            "total_words": extracted_content["metadata"]["total_words"],
            "has_tables": len(extracted_content["tables"]) > 0,
            "has_headers": extracted_content["metadata"]["has_headers"]
        }

        # Analyze with AI
        analysis_result = await ai_template_analyzer.analyze_document_for_template(
            clean_text,
            openai_api_key,
            metadata
        )

        print(f"[AI ANALYZE] Analysis complete - found {len(analysis_result.variables)} variables")

        # Log the analysis
        db.add(ActivityLog(
            action="template.ai_analyze",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template_id,
            message=f"AI analysis performed on template '{template.name}'",
            log_metadata={
                "variables_found": len(analysis_result.variables),
                "confidence_score": analysis_result.confidence_score,
                "suggested_type": analysis_result.template_type
            }
        ))
        db.commit()

        # Return structured response
        return {
            "template_id": template_id,
            "analysis": {
                "template_type": analysis_result.template_type,
                "document_category": analysis_result.document_category,
                "confidence_score": analysis_result.confidence_score,
                "placeholder_format": analysis_result.placeholder_format,
                "summary": analysis_result.summary,
                "variables": [
                    {
                        "name": var.name,
                        "original_text": var.original_text,
                        "suggested_placeholder": var.suggested_placeholder,
                        "description": var.description,
                        "type": var.type,
                        "context": var.context,
                        "confidence": var.confidence,
                        "start_position": var.start_position,
                        "end_position": var.end_position
                    }
                    for var in analysis_result.variables
                ],
                "text_improvements": [
                    {
                        "location": improvement.location,
                        "original_text": improvement.original_text,
                        "improved_text": improvement.improved_text,
                        "improvement_type": improvement.improvement_type
                    }
                    for improvement in analysis_result.text_improvements
                ],
                "suggestions": analysis_result.suggestions
            },
            "extracted_content": {
                "total_paragraphs": extracted_content["metadata"]["total_paragraphs"],
                "total_tables": extracted_content["metadata"]["total_tables"],
                "total_words": extracted_content["metadata"]["total_words"],
                "has_headers": extracted_content["metadata"]["has_headers"],
                "has_footers": extracted_content["metadata"]["has_footers"]
            },
            "status": "completed"
        }

    except Exception as e:
        print(f"[AI ANALYZE ERROR] {str(e)}")

        # Log the error
        db.add(ActivityLog(
            action="template.ai_analyze_error",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template_id,
            message=f"AI analysis failed for template '{template.name}': {str(e)}",
            log_metadata={"error": str(e)}
        ))
        db.commit()

        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

@router.post("/ai-apply/{template_id}")
async def apply_ai_suggestions(
    template_id: int,
    request: Request,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """
    Apply accepted AI suggestions to create an improved template
    """
    print(f"[AI APPLY] Applying AI suggestions to template {template_id}")

    try:
        # Parse request body
        body = await request.json()

        # Handle both old format (accepted_variables) and new format (variables/improvements)
        if "accepted_variables" in body:
            # Legacy format - just variables
            accepted_variables = body.get("accepted_variables", [])
            accepted_improvements = []
        else:
            # New format - both variables and improvements
            accepted_variables = body.get("variables", [])
            accepted_improvements = body.get("improvements", [])

        new_template_name = body.get("new_template_name")
        openai_api_key = body.get("openai_api_key")

        if not accepted_variables and not accepted_improvements:
            raise HTTPException(status_code=400, detail="No variables or improvements to apply")

        template = db.query(Template).filter(Template.id == template_id).first()
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")

        # Get the current document content
        if not file_storage.file_exists(template_id):
            raise HTTPException(status_code=400, detail="No DOCX file available")

        file_content = file_storage.get_docx_content(template_id)

        # Create replacements dictionary for formatting-preserving replacement
        all_replacements = {}

        # Add variable replacements
        for var_data in accepted_variables:
            original_text = var_data["original_text"]
            placeholder = var_data["suggested_placeholder"]
            all_replacements[original_text] = placeholder

        # Add text improvements
        for improvement_data in accepted_improvements:
            original_text = improvement_data["original_text"]
            improved_text = improvement_data["improved_text"]
            all_replacements[original_text] = improved_text

        print(f"[AI APPLY] Applying {len(accepted_variables)} variables and {len(accepted_improvements)} improvements")
        for original, replacement in all_replacements.items():
            print(f"[AI APPLY]   '{original}' -> '{replacement}'")

        # Apply all replacements while preserving ALL formatting
        improved_file_content = document_text_improver.apply_all_replacements_to_document(
            file_content, all_replacements
        )

        # Generate variable definitions for template metadata
        from app.services.ai_template_analyzer import TemplateVariable
        variables = []
        for var_data in accepted_variables:
            variable = TemplateVariable(
                name=var_data["name"],
                original_text=var_data["original_text"],
                suggested_placeholder=var_data["suggested_placeholder"],
                description=var_data["description"],
                type=var_data["type"],
                context=var_data["context"],
                confidence=var_data["confidence"],
                start_position=var_data.get("start_position"),
                end_position=var_data.get("end_position")
            )
            variables.append(variable)

        variable_definitions = ai_template_analyzer.generate_variable_definitions(variables)

        # Create new template or update existing
        if new_template_name:
            # Create new template
            new_template_data = {
                "name": new_template_name,
                "description": f"AI-enhanced template from '{template.name}' - formatting preserved",
                "template_type": template.template_type,
                "content": {
                    "type": "ai_enhanced",
                    "original_template_id": template_id,
                    "variables_applied": len(accepted_variables),
                    "improvements_applied": len(accepted_improvements)
                },
                "variables": variable_definitions,
                "is_ai_enhanced": True,
                "status": "saved",
                "created_by": current_user.id
            }

            new_template = Template(**new_template_data)
            db.add(new_template)
            db.commit()
            db.refresh(new_template)

            # Save the improved document (preserves all formatting)
            file_metadata = file_storage.save_docx_content(improved_file_content, new_template.id)
            new_template.file_path = f"uploads/templates/template_{new_template.id}.docx"
            new_template.file_name = f"{new_template.name}.docx"
            new_template.file_size = len(improved_file_content)
            new_template.mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            db.commit()

            result_template_id = new_template.id
            action = "create"
        else:
            # Update existing template
            file_metadata = file_storage.save_docx_content(improved_file_content, template_id)
            template.variables = variable_definitions
            template.is_ai_enhanced = True
            template.updated_at = datetime.utcnow()

            # Update content to include AI modifications
            if isinstance(template.content, dict):
                template.content.update({
                    "ai_enhanced": True,
                    "variables_applied": len(accepted_variables),
                    "improvements_applied": len(accepted_improvements)
                })

            db.commit()
            result_template_id = template_id
            action = "update"

        # Log the successful application
        db.add(ActivityLog(
            action=f"template.ai_apply_{action}",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=result_template_id,
            message=f"Applied {len(accepted_variables)} variables and {len(accepted_improvements)} improvements to template (formatting preserved)",
            log_metadata={
                "original_template_id": template_id,
                "variables_applied": len(all_replacements)
            }
        ))
        db.commit()

        return {
            "status": "success",
            "action": action,
            "template_id": result_template_id,
            "variables_applied": len(all_replacements),
            "message": f"Successfully applied {len(accepted_variables)} variables and {len(accepted_improvements)} improvements while preserving all formatting"
        }

    except Exception as e:
        print(f"[AI APPLY ERROR] {str(e)}")

        # Log the error
        db.add(ActivityLog(
            action="template.ai_apply_error",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template_id,
            message=f"Failed to apply AI suggestions: {str(e)}",
            log_metadata={"error": str(e)}
        ))
        db.commit()

        raise HTTPException(status_code=500, detail=f"Failed to apply suggestions: {str(e)}")

@router.get("/ai-extract-text/{template_id}")
async def extract_template_text(
    template_id: int,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """
    Extract clean text from a template for AI analysis preview
    """
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    try:
        if not file_storage.file_exists(template_id):
            raise HTTPException(status_code=400, detail="No DOCX file available")

        # Get the DOCX file content and extract text
        file_content = file_storage.get_docx_content(template_id)
        extracted_content = document_extractor.extract_from_content(file_content)

        # Also extract potential variables for preview
        clean_text = document_extractor.get_clean_text_for_ai(extracted_content)
        potential_variables = document_extractor.extract_potential_variables(clean_text)

        return {
            "template_id": template_id,
            "extracted_content": extracted_content,
            "clean_text": clean_text,
            "potential_variables": potential_variables,
            "metadata": extracted_content["metadata"]
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(e)}")

# Text Improvement Endpoints

@router.post("/improve-text/{template_id}")
async def improve_template_text(
    template_id: int,
    openai_api_key: str = Form(...),
    improvement_type: str = Form(default="grammar_clarity"),
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """
    Improve text content while preserving all formatting, layout, and structure
    """
    print(f"[TEXT IMPROVE] Starting text improvement for template {template_id}")

    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    try:
        # Check if DOCX file exists
        if not template.file_path or not file_storage.file_exists(template_id):
            raise HTTPException(status_code=400, detail="No DOCX file available for text improvement")

        # Get the DOCX file content
        file_content = file_storage.get_docx_content(template_id)

        # Extract text segments
        segments = document_text_improver.extract_text_segments(file_content)
        print(f"[TEXT IMPROVE] Extracted {len(segments)} text segments")

        if not segments:
            raise HTTPException(status_code=400, detail="No text content found to improve")

        # Improve text with AI
        improvement_result = await document_text_improver.improve_text_with_ai(
            segments,
            openai_api_key,
            improvement_type
        )

        print(f"[TEXT IMPROVE] Improved {improvement_result.improved_segments} segments")

        # Generate preview
        preview = document_text_improver.get_improvement_preview(improvement_result.segments)

        # Log the improvement
        db.add(ActivityLog(
            action="template.text_improve",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template_id,
            message=f"Text improvement performed on template '{template.name}'",
            log_metadata={
                "improvement_type": improvement_type,
                "total_segments": improvement_result.total_segments,
                "improved_segments": improvement_result.improved_segments,
                "preserved_placeholders": len(improvement_result.preserved_placeholders)
            }
        ))
        db.commit()

        return {
            "template_id": template_id,
            "improvement_result": {
                "total_segments": improvement_result.total_segments,
                "improved_segments": improvement_result.improved_segments,
                "preserved_placeholders": improvement_result.preserved_placeholders,
                "summary": improvement_result.summary
            },
            "preview": preview,
            "status": "completed"
        }

    except Exception as e:
        print(f"[TEXT IMPROVE ERROR] {str(e)}")

        # Log the error
        db.add(ActivityLog(
            action="template.text_improve_error",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template_id,
            message=f"Text improvement failed for template '{template.name}': {str(e)}",
            log_metadata={"error": str(e)}
        ))
        db.commit()

        raise HTTPException(status_code=500, detail=f"Text improvement failed: {str(e)}")

@router.post("/apply-text-improvements/{template_id}")
async def apply_text_improvements(
    template_id: int,
    request: Request,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """
    Apply text improvements to the document while preserving all formatting
    """
    print(f"[APPLY TEXT] Applying text improvements to template {template_id}")

    try:
        # Parse request body
        body = await request.json()
        improved_segments_data = body.get("improved_segments", [])
        create_new_template = body.get("create_new_template", True)
        new_template_name = body.get("new_template_name")

        if not improved_segments_data:
            raise HTTPException(status_code=400, detail="No improved segments to apply")

        template = db.query(Template).filter(Template.id == template_id).first()
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")

        # Get the current document content
        if not file_storage.file_exists(template_id):
            raise HTTPException(status_code=400, detail="No DOCX file available")

        file_content = file_storage.get_docx_content(template_id)

        # Convert improved segments data back to TextSegment objects
        from app.services.text_improver import TextSegment
        improved_segments = []
        for segment_data in improved_segments_data:
            segment = TextSegment(
                original_text=segment_data["original_text"],
                location_type=segment_data["location_type"],
                location_info=segment_data["location_info"],
                improved_text=segment_data.get("improved_text")
            )
            improved_segments.append(segment)

        # Apply improvements to document
        improved_file_content = document_text_improver.apply_improvements_to_document(
            file_content, improved_segments
        )

        if create_new_template:
            # Create new template with improved text
            new_template_data = {
                "name": new_template_name or f"{template.name} (Text Improved)",
                "description": f"Text-improved version of '{template.name}' - formatting preserved",
                "template_type": template.template_type,
                "content": {
                    "type": "text_improved",
                    "original_template_id": template_id,
                    "improvement_applied": True
                },
                "variables": template.variables,
                "is_ai_enhanced": True,
                "status": "saved",
                "created_by": current_user.id
            }

            new_template = Template(**new_template_data)
            db.add(new_template)
            db.commit()
            db.refresh(new_template)

            # Save the improved document
            file_metadata = file_storage.save_docx_content(improved_file_content, new_template.id)
            new_template.file_path = f"uploads/templates/template_{new_template.id}.docx"
            new_template.file_name = f"{new_template.name}.docx"
            new_template.file_size = len(improved_file_content)
            new_template.mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            db.commit()

            result_template_id = new_template.id
            action = "create"
        else:
            # Update existing template
            file_metadata = file_storage.save_docx_content(improved_file_content, template_id)
            template.updated_at = datetime.utcnow()

            # Update content metadata
            if isinstance(template.content, dict):
                template.content.update({
                    "text_improved": True,
                    "improvement_applied": True
                })

            db.commit()
            result_template_id = template_id
            action = "update"

        # Log the successful application
        db.add(ActivityLog(
            action=f"template.text_improve_apply_{action}",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=result_template_id,
            message=f"Applied text improvements to template (formatting preserved)",
            log_metadata={
                "original_template_id": template_id,
                "segments_improved": len([s for s in improved_segments if s.improved_text])
            }
        ))
        db.commit()

        return {
            "status": "success",
            "action": action,
            "template_id": result_template_id,
            "improvements_applied": len([s for s in improved_segments if s.improved_text]),
            "message": f"Successfully applied text improvements while preserving all formatting"
        }

    except Exception as e:
        print(f"[APPLY TEXT ERROR] {str(e)}")

        # Log the error
        db.add(ActivityLog(
            action="template.text_improve_apply_error",
            actor_user_id=current_user.id,
            target_type="template",
            target_id=template_id,
            message=f"Failed to apply text improvements: {str(e)}",
            log_metadata={"error": str(e)}
        ))
        db.commit()

        raise HTTPException(status_code=500, detail=f"Failed to apply improvements: {str(e)}")


@router.get("/{template_id}/placeholders")
def extract_template_placeholders(
    template_id: int,
    current_user: UserOut = Depends(require_admin_or_superadmin),
    db: Session = Depends(get_db)
):
    """Extract all placeholders from a template DOCX file"""
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    # Check if template has a file
    if not template.file_path or not file_storage.file_exists(template_id):
        raise HTTPException(status_code=400, detail="Template file not found")

    try:
        # Get template content
        template_content = file_storage.get_docx_content(template_id)

        # Extract placeholders
        placeholders = quotation_filler.extract_all_placeholders(template_content)

        return {
            "template_id": template_id,
            "template_name": template.name,
            "placeholders": placeholders,
            "total": len(placeholders)
        }

    except Exception as e:
        print(f"[ERROR] Failed to extract placeholders: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract placeholders: {str(e)}")

